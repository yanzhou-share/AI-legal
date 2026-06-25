"""
合同感知模块 - ContractPerception
负责读取 Word 文档，将合同按段落和表格单元格切片
"""
from dataclasses import dataclass
from typing import List, Tuple
from pathlib import Path

from docx import Document
from docx.table import Table, _Cell


@dataclass
class TextBlock:
    """文本块数据结构"""
    block_id: str
    text: str
    block_type: str  # 'paragraph' 或 'table_cell'
    index: int  # 在文档中的位置索引


class ContractPerception:
    """合同感知模块 - 负责文档解析和文本块切片"""

    def __init__(self):
        self._blocks: List[TextBlock] = []
        self._block_counter = 0

    def _generate_block_id(self) -> str:
        """生成唯一的 block_id"""
        self._block_counter += 1
        return f"block_{self._block_counter:04d}"

    def scan_contract(self, file_path: str) -> List[TextBlock]:
        """
        读取合同并切片为独立文本块
        
        Args:
            file_path: Word 文档路径
            
        Returns:
            文本块列表，每个块包含 block_id、文本内容和类型
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"合同文件不存在: {file_path}")

        self._blocks.clear()
        self._block_counter = 0

        doc = Document(file_path)
        self._process_body(doc)
        self._process_tables(doc.tables)

        return self._blocks

    def _process_body(self, doc: Document) -> None:
        """
        处理文档主体内容（段落）
        
        Args:
            doc: 文档对象
        """
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                block = TextBlock(
                    block_id=self._generate_block_id(),
                    text=text,
                    block_type='paragraph',
                    index=idx
                )
                self._blocks.append(block)

    def _process_tables(self, tables: List[Table]) -> None:
        """
        处理表格，为每个段落创建文本块
        
        Args:
            tables: 表格列表
        """
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # 处理单元格中的每个段落
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            block = TextBlock(
                                block_id=self._generate_block_id(),
                                text=text,
                                block_type='table_cell',
                                index=len(self._blocks)
                            )
                            self._blocks.append(block)

    def get_block_by_id(self, block_id: str) -> TextBlock:
        """根据 block_id 获取文本块"""
        for block in self._blocks:
            if block.block_id == block_id:
                return block
        raise ValueError(f"未找到 block_id: {block_id}")

    def get_all_blocks(self) -> List[TextBlock]:
        """获取所有文本块"""
        return self._blocks
