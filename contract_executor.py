"""
无损回填与执行模块 - ContractExecutor
负责将修改后的条款精准回填到 Word 文档中，保持原有格式
"""
import re
from typing import Dict, List, Tuple
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.shared import RGBColor


class ContractExecutor:
    """无损回填与执行模块 - 负责文档的精准修改和保存"""

    def execute_modification(
        self, 
        file_path: str, 
        output_path: str, 
        modification_map: Dict[str, str]
    ) -> bool:
        """
        执行文档修改，将修改后的文本回填到原文档
        
        Args:
            file_path: 原文档路径
            output_path: 输出文档路径
            modification_map: 修改映射 {block_id: revised_text}
            
        Returns:
            是否修改成功
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"原文档不存在: {file_path}")

        if not modification_map:
            print("[ContractExecutor] 没有需要修改的内容")
            return False

        doc = Document(file_path)
        
        # 构建段落到 block_id 的映射（与 perception 保持一致）
        para_block_map = self._build_paragraph_block_map(doc)
        
        # 执行修改
        for block_id, new_text in modification_map.items():
            if block_id in para_block_map:
                para = para_block_map[block_id]
                self._modify_paragraph(para, new_text, mark_red=True)

        # 确保输出路径以 .docx 结尾
        if not output_path.endswith('.docx'):
            output_path = output_path.rsplit('.', 1)[0] + '.docx'

        # 保存文档
        output_dir = Path(output_path).parent
        if output_dir and not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)

        doc.save(output_path)
        print(f"[ContractExecutor] 文档已保存至: {output_path}")
        return True

    def _build_paragraph_block_map(self, doc: Document) -> Dict[str, Paragraph]:
        """
        构建段落到 block_id 的映射
        
        Args:
            doc: 文档对象
            
        Returns:
            {block_id: paragraph} 映射
        """
        para_map = {}
        block_counter = 0
        
        # 遍历所有段落（包括表格中的段落）
        # 这与 ContractPerception.scan_contract 的逻辑保持一致
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                block_counter += 1
                block_id = f"block_{block_counter:04d}"
                para_map[block_id] = para
        
        # 处理表格中的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            block_counter += 1
                            block_id = f"block_{block_counter:04d}"
                            para_map[block_id] = para
        
        return para_map

    def _modify_paragraph(self, paragraph: Paragraph, new_text: str, mark_red: bool = True) -> None:
        """
        无损修改段落内容，可选标红显示
        
        Args:
            paragraph: 段落对象
            new_text: 新文本
            mark_red: 是否标红显示修改内容
        """
        # 清理文本
        clean_text = new_text.replace('\n', ' ').replace('\r', '')
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()

        # 保存段落格式
        para_format = {
            'alignment': paragraph.paragraph_format.alignment,
            'first_line_indent': paragraph.paragraph_format.first_line_indent,
            'left_indent': paragraph.paragraph_format.left_indent,
            'right_indent': paragraph.paragraph_format.right_indent,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after,
            'line_spacing': paragraph.paragraph_format.line_spacing,
        }

        # 保存第一个 Run 的字体格式
        base_font_name = None
        base_font_size = None
        base_font_bold = None
        base_font_italic = None
        base_font_underline = None
        
        if paragraph.runs:
            base_run = paragraph.runs[0]
            base_font_name = base_run.font.name
            base_font_size = base_run.font.size
            base_font_bold = base_run.font.bold
            base_font_italic = base_run.font.italic
            base_font_underline = base_run.font.underline

        # 清空所有 Run 的文本
        for run in paragraph.runs:
            run.text = ""

        # 设置新文本到第一个 Run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run.text = clean_text
            
            # 恢复字体格式
            if base_font_name:
                first_run.font.name = base_font_name
            if base_font_size:
                first_run.font.size = base_font_size
            if base_font_bold is not None:
                first_run.font.bold = base_font_bold
            if base_font_italic is not None:
                first_run.font.italic = base_font_italic
            if base_font_underline is not None:
                first_run.font.underline = base_font_underline
            
            # 标红显示修改内容
            if mark_red:
                first_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
        else:
            # 没有 Run，创建一个新的
            run = paragraph.add_run(clean_text)
            if base_font_name:
                run.font.name = base_font_name
            if base_font_size:
                run.font.size = base_font_size
            if mark_red:
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # 恢复段落格式
        paragraph.paragraph_format.alignment = para_format['alignment']
        paragraph.paragraph_format.first_line_indent = para_format['first_line_indent']
        paragraph.paragraph_format.left_indent = para_format['left_indent']
        paragraph.paragraph_format.right_indent = para_format['right_indent']
        paragraph.paragraph_format.space_before = para_format['space_before']
        paragraph.paragraph_format.space_after = para_format['space_after']
        paragraph.paragraph_format.line_spacing = para_format['line_spacing']

    def _modify_cell(self, cell, new_text: str) -> None:
        """
        无损修改表格单元格内容
        
        Args:
            cell: 单元格对象
            new_text: 新文本
        """
        # 获取单元格中的第一个段落
        if cell.paragraphs:
            para = cell.paragraphs[0]
            self._modify_paragraph(para, new_text, mark_red=True)

            # 清空其他段落
            for extra_para in cell.paragraphs[1:]:
                for run in extra_para.runs:
                    run.text = ""

    def preview_modification(
        self, 
        file_path: str, 
        modification_map: Dict[str, str]
    ) -> List[Dict]:
        """
        预览修改内容（不实际修改文件）
        
        Args:
            file_path: 文档路径
            modification_map: 修改映射
            
        Returns:
            修改预览列表
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"文档不存在: {file_path}")

        doc = Document(file_path)
        previews = []
        
        # 构建映射
        para_map = self._build_paragraph_block_map(doc)
        
        for block_id, new_text in modification_map.items():
            if block_id in para_map:
                para = para_map[block_id]
                previews.append({
                    'block_id': block_id,
                    'original': para.text.strip(),
                    'revised': new_text[:100] + '...' if len(new_text) > 100 else new_text
                })

        return previews
