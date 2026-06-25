"""
智能法务合同修改与润色智能体 - 主程序
完整的测试工作流演示
"""
import sys
import io
from pathlib import Path

# 设置标准输出编码为 UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from contract_perception import ContractPerception
from legal_brain import LegalBrain
from security_guard import SecurityGuard
from contract_executor import ContractExecutor


def create_test_contract(file_path: str) -> None:
    """
    创建测试用合同文档
    
    Args:
        file_path: 输出文件路径
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '仿宋_GB2312'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('技术服务合同', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 合同编号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('合同编号：TECH-2024-001')
    run.font.size = Pt(10)

    # 甲方信息
    doc.add_paragraph('甲方：北京科技创新有限公司')
    doc.add_paragraph('乙方：上海软件开发工作室')

    # 合同正文
    doc.add_paragraph('')
    doc.add_paragraph('第一条 服务内容')
    doc.add_paragraph(
        '乙方为甲方提供企业管理系统开发服务，包括需求分析、系统设计、'
        '编码实现、测试部署及后期维护，项目周期为6个月。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('第二条 服务费用')
    doc.add_paragraph(
        '本合同总金额为人民币壹佰伍拾万元整（¥1,500,000.00），'
        '甲方应于合同签订后5个工作日内支付30%预付款，项目验收合格后支付剩余70%。'
    )

    # 不平等条款 - 违约金条款
    doc.add_paragraph('')
    doc.add_paragraph('第三条 违约责任')
    p = doc.add_paragraph(
        '若乙方未能按时完成项目交付，每延迟一日，应向甲方支付合同总金额10%的违约金；'
        '若延迟超过30日，甲方有权解除合同并要求乙方双倍返还已支付款项。'
        '甲方延迟付款的，乙方不得停止工作，且不承担任何违约责任。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('第四条 知识产权')
    doc.add_paragraph(
        '本项目开发的所有软件、源代码、技术文档的知识产权归甲方所有。'
        '乙方不得将上述成果用于其他商业用途。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('第五条 保密条款')
    doc.add_paragraph(
        '双方对本合同内容及履行过程中知悉的对方商业秘密负有保密义务，'
        '保密期限为合同终止后3年。'
    )

    doc.add_paragraph('')
    doc.add_paragraph('第六条 争议解决')
    doc.add_paragraph(
        '本合同在履行过程中发生争议，双方应友好协商解决；'
        '协商不成的，任何一方均可向甲方所在地人民法院提起诉讼。'
    )

    # 签章区域
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph('甲方（盖章）：北京科技创新有限公司')
    p = doc.add_paragraph('法定代表人：张三')
    p = doc.add_paragraph('日期：2024年1月15日')

    doc.add_paragraph('')
    p = doc.add_paragraph('乙方（盖章）：上海软件开发工作室')
    p = doc.add_paragraph('法定代表人：李四')
    p = doc.add_paragraph('日期：2024年1月15日')

    doc.save(file_path)
    print(f"[测试] 已创建测试合同: {file_path}")


def run_contract_agent(
    file_path: str,
    output_path: str,
    user_instruction: str,
    api_key: str = "test-key",
    base_url: str = "https://api.openai.com/v1",
    model: str = "gpt-4"
) -> bool:
    """
    运行完整的合同修改智能体工作流
    
    Args:
        file_path: 输入合同文件路径
        output_path: 输出合同文件路径
        user_instruction: 用户修改指令
        api_key: API 密钥
        base_url: API 基础 URL
        model: 模型名称
        
    Returns:
        是否成功
    """
    print("=" * 60)
    print("智能法务合同修改与润色智能体 - 开始执行")
    print("=" * 60)

    # 步骤1: 合同感知 - 扫描和切片
    print("\n[步骤1] 合同感知 - 扫描文档...")
    perception = ContractPerception()
    blocks = perception.scan_contract(file_path)
    print(f"  - 发现 {len(blocks)} 个文本块")

    # 步骤2: 安全审计器初始化
    print("\n[步骤2] 初始化安全审计模块...")
    security = SecurityGuard()

    # 步骤3: 法务大模型决策
    print("\n[步骤3] 法务大模型审查与润色...")
    brain = LegalBrain(api_key=api_key, base_url=base_url, model=model)

    modification_map = {}
    for block in blocks:
        print(f"  - 审查块 {block.block_id}: {block.text[:30]}...")

        result = brain.review_and_rewrite(user_instruction, block.text)
        if result is None:
            print(f"    [警告] 块 {block.block_id} 审查失败，跳过")
            continue

        # 安全审计
        audit_result, audit_msg = security.audit_json_response({
            'revised_text': result.revised_text,
            'risk_analysis': result.risk_analysis
        })

        if not audit_result:
            print(f"    [警告] 块 {block.block_id} 安全审计失败: {audit_msg}")
            continue

        if result.has_risk and result.revised_text != block.text:
            modification_map[block.block_id] = result.revised_text
            print(f"    [发现风险] {result.risk_analysis[:50]}...")
        else:
            print(f"    [无风险] 保持原文")

    print(f"\n  - 共发现 {len(modification_map)} 处需要修改的条款")

    # 步骤4: 预览修改
    if modification_map:
        print("\n[步骤4] 修改预览:")
        executor = ContractExecutor()
        previews = executor.preview_modification(file_path, modification_map)
        for preview in previews:
            print(f"\n  [{preview['block_id']}]")
            print(f"    原文: {preview['original'][:80]}...")
            print(f"    修改: {preview['revised'][:80]}...")

    # 步骤5: 无损回填
    print("\n[步骤5] 无损回填与保存...")
    executor = ContractExecutor()
    success = executor.execute_modification(file_path, output_path, modification_map)

    if success:
        print("\n" + "=" * 60)
        print("智能法务合同修改与润色智能体 - 执行完成")
        print("=" * 60)
        print(f"输入文件: {file_path}")
        print(f"输出文件: {output_path}")
        print(f"修改条数: {len(modification_map)}")
    else:
        print("\n[完成] 合同无需修改或修改失败")

    return success


def main():
    """主函数 - 完整测试工作流"""
    print("\n" + "=" * 60)
    print("智能法务合同修改与润色智能体 - 测试工作流")
    print("=" * 60 + "\n")

    # 配置
    test_contract_path = "test_contract.docx"
    modified_contract_path = "modified_contract.docx"
    user_instruction = "把里面对乙方不利的单方面高额违约金条款修改为双方平等的对等条款"

    # 步骤1: 创建测试合同
    print("[准备] 创建测试合同文档...")
    create_test_contract(test_contract_path)

    # 步骤2: 运行智能体（使用模拟配置）
    # 注意：实际使用时需要提供真实的 API 密钥和基础 URL
    print("\n[运行] 启动智能体工作流...")
    print("  [提示] 当前使用模拟配置，实际运行需要配置 API 密钥")
    print(f"  [指令] {user_instruction}\n")

    # 这里使用模拟数据演示完整流程
    # 实际运行时取消下面的注释并配置真实的 API 信息
    # run_contract_agent(
    #     file_path=test_contract_path,
    #     output_path=modified_contract_path,
    #     user_instruction=user_instruction,
    #     api_key="your-api-key",
    #     base_url="https://api.openai.com/v1",
    #     model="gpt-4"
    # )

    # 模拟演示流程
    print("[演示] 模拟执行流程...\n")

    # 1. 合同感知
    perception = ContractPerception()
    blocks = perception.scan_contract(test_contract_path)
    print(f"[感知] 发现 {len(blocks)} 个文本块")
    for block in blocks:
        print(f"  - {block.block_id}: {block.text[:60]}...")

    # 2. 安全审计
    security = SecurityGuard()
    test_safe = security.audit_json_response({
        'revised_text': '这是一段安全的合同文本',
        'risk_analysis': '无风险'
    })
    print(f"\n[安全] 审计测试: {'通过' if test_safe[0] else '失败'}")

    # 3. 模拟 LLM 返回（实际使用时由 LegalBrain 类处理）
    print("\n[决策] 模拟法务审查结果...")
    simulated_modifications = {
        'block_0006': (
            '若乙方未能按时完成项目交付，每延迟一日，应向甲方支付合同总金额1%的违约金；'
            '若延迟超过30日，甲方有权解除合同并要求乙方返还已支付款项。'
            '甲方延迟付款超过15日的，乙方有权暂停工作，'
            '甲方延迟付款超过30日的，乙方有权解除合同并要求甲方支付已完成工作的费用。'
        )
    }

    # 4. 预览修改
    executor = ContractExecutor()
    previews = executor.preview_modification(test_contract_path, simulated_modifications)
    print("\n[预览] 修改内容:")
    for p in previews:
        print(f"\n  [{p['block_id']}]")
        print(f"    原文: {p['original'][:80]}...")
        print(f"    修改: {p['revised'][:80]}...")

    # 5. 执行修改
    print("\n[执行] 无损回填修改...")
    success = executor.execute_modification(
        test_contract_path, 
        modified_contract_path, 
        simulated_modifications
    )

    print("\n" + "=" * 60)
    if success:
        print("测试工作流执行成功！")
        print(f"已生成修改后的合同: {modified_contract_path}")
    else:
        print("测试工作流执行完成（模拟模式）")
    print("=" * 60)


if __name__ == "__main__":
    main()
